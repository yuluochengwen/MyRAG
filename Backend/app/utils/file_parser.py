"""文件解析器 - 将不同格式的文件转换为文本"""
import os
import re
from pathlib import Path
from typing import Any, Optional, List
from abc import ABC, abstractmethod

from app.utils.logger import get_logger

logger = get_logger(__name__)


class MarkdownNormalizer:
    """Markdown文本规范化工具"""

    @staticmethod
    def normalize(text: str) -> str:
        if not text:
            return ""

        normalized = text.replace('\r\n', '\n').replace('\r', '\n')
        normalized = normalized.replace('\u00a0', ' ')
        normalized = re.sub(r"[\t\f\v]+", " ", normalized)
        normalized = re.sub(r"\n{3,}", "\n\n", normalized)
        normalized = re.sub(r"[ ]{2,}", " ", normalized)
        return normalized.strip()

    @staticmethod
    def looks_like_markdown(text: str) -> bool:
        if not text:
            return False
        patterns = [r"^#{1,6}\s", r"^[-*]\s", r"^\d+\.\s", r"^\|.*\|", r"^```", r"\n##\s"]
        return any(re.search(pattern, text, flags=re.MULTILINE) for pattern in patterns)

    @staticmethod
    def infer_markdown_structure(text: str) -> str:
        """为普通纯文本补齐部分 Markdown 结构（温和规则）。"""
        if not text:
            return ""

        lines = [line.strip() for line in text.split('\n')]
        result: List[str] = []

        heading_patterns = [
            r"^第[一二三四五六七八九十百0-9]+章",
            r"^[一二三四五六七八九十]+、",
            r"^\(?[一二三四五六七八九十]+\)",
            r"^\d+(?:\.\d+){0,2}\s+"
        ]

        for line in lines:
            if not line:
                result.append("")
                continue

            if any(re.match(pattern, line) for pattern in heading_patterns):
                result.append(f"## {line}")
                continue

            if line.startswith('- ') or re.match(r"^\d+\.\s", line):
                result.append(line)
                continue

            result.append(line)

        return MarkdownNormalizer.normalize('\n'.join(result))


class BaseParser(ABC):
    """文件解析器基类"""
    
    @abstractmethod
    def parse(self, file_path: str) -> str:
        """
        解析文件为文本
        
        Args:
            file_path: 文件路径
            
        Returns:
            解析后的文本内容
        """
        pass


class TxtParser(BaseParser):
    """TXT文件解析器"""
    
    def parse(self, file_path: str) -> str:
        """解析TXT文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            for encoding in ['gbk', 'gb2312', 'latin-1']:
                try:
                    with open(file_path, 'r', encoding=encoding) as f:
                        return f.read()
                except (UnicodeDecodeError, LookupError):
                    continue
            
            raise ValueError(f"无法解析文件编码: {file_path}")


class MarkdownParser(BaseParser):
    """Markdown文件解析器"""
    
    def parse(self, file_path: str) -> str:
        """解析Markdown文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except Exception as e:
            logger.error(f"解析Markdown文件失败: {file_path}, {str(e)}")
            raise


class PDFParser(BaseParser):
    """PDF文件解析器 - 增强版，保留段落结构"""
    
    def parse(self, file_path: str) -> str:
        """解析PDF文件并优先转换为Markdown"""
        # 优先尝试高质量 PDF -> Markdown 转换
        markdown_text = self._parse_with_pymupdf4llm(file_path)
        if markdown_text:
            return markdown_text

        # 二级回退：使用PyMuPDF直接提取文本层（对部分中文PDF更稳）
        pymupdf_text = self._parse_with_pymupdf_text(file_path)
        if pymupdf_text:
            return pymupdf_text

        # 回退到现有 PyPDF2 解析
        try:
            from PyPDF2 import PdfReader
            
            reader = PdfReader(file_path)
            pages = []
            
            for page_num, page in enumerate(reader.pages, 1):
                page_text = page.extract_text()
                if page_text:
                    # 清理页面文本：移除过多的空行，但保留段落分隔
                    lines = page_text.split('\n')
                    cleaned_lines = []
                    prev_empty = False
                    
                    for line in lines:
                        line = line.strip()
                        if line:
                            cleaned_lines.append(line)
                            prev_empty = False
                        elif not prev_empty:
                            # 保留一个空行作为段落分隔
                            cleaned_lines.append('')
                            prev_empty = True
                    
                    # 用双换行连接段落（为递归分割提供明确的边界）
                    page_content = '\n'.join(cleaned_lines)
                    
                    # 添加页码标记（可选，帮助追踪来源）
                    pages.append(f"[第{page_num}页]\n{page_content}")
            
            # 页面间用双换行分隔（段落级别的边界）
            plain_text = '\n\n'.join(pages)
            structured = MarkdownNormalizer.infer_markdown_structure(plain_text)
            if structured:
                return structured
            raise ValueError("PDF文本层为空，可能为扫描件或加密文档")
            
        except Exception as e:
            logger.error(f"解析PDF文件失败: {file_path}, {str(e)}")
            raise

    def _parse_with_pymupdf4llm(self, file_path: str) -> Optional[str]:
        try:
            import pymupdf4llm

            markdown_text = pymupdf4llm.to_markdown(file_path)
            if markdown_text and len(markdown_text.strip()) > 0:
                logger.info(f"PDF使用pymupdf4llm转换Markdown成功: {file_path}")
                return MarkdownNormalizer.normalize(markdown_text)
        except Exception as error:
            logger.warning(f"pymupdf4llm不可用或转换失败，回退PyPDF2: {str(error)}")
        return None

    def _parse_with_pymupdf_text(self, file_path: str) -> Optional[str]:
        """使用PyMuPDF提取PDF文本层，避免仅依赖OCR。"""
        try:
            import fitz

            pages: List[str] = []
            with fitz.open(file_path) as document:
                for page_num, page in enumerate(document, 1):
                    page_text = self._extract_text_from_page(page)
                    if page_text and page_text.strip():
                        cleaned = MarkdownNormalizer.normalize(page_text)
                        if cleaned:
                            pages.append(f"[第{page_num}页]\n{cleaned}")

            if not pages:
                return None

            merged = "\n\n".join(pages)
            normalized = MarkdownNormalizer.infer_markdown_structure(merged)
            if normalized:
                logger.info(f"PDF使用PyMuPDF文本层提取成功: {file_path}")
                return normalized
            return None
        except Exception as error:
            logger.warning(f"PyMuPDF文本层提取失败，回退PyPDF2: {str(error)}")
            return None

    def _extract_text_from_page(self, page: Any) -> str:
        """多策略提取单页文本，兼容不同PDF编码/布局。"""
        # 1) 直接文本提取
        text = page.get_text("text") or ""
        if text.strip():
            return text

        # 2) 块级提取（部分文档text为空但blocks可用）
        try:
            blocks = page.get_text("blocks") or []
            block_texts = [str(item[4]).strip() for item in blocks if isinstance(item, (list, tuple)) and len(item) > 4 and str(item[4]).strip()]
            if block_texts:
                return "\n".join(block_texts)
        except Exception:
            pass

        # 3) 词级提取（进一步兜底）
        try:
            words = page.get_text("words") or []
            if words:
                words = sorted(words, key=lambda row: (float(row[1]), float(row[0])))
                lines: List[str] = []
                current_line_y: Optional[float] = None
                current_words: List[str] = []

                for row in words:
                    if len(row) < 5:
                        continue
                    y = float(row[1])
                    token = str(row[4]).strip()
                    if not token:
                        continue
                    if current_line_y is None or abs(y - current_line_y) <= 3.0:
                        current_words.append(token)
                        current_line_y = y if current_line_y is None else current_line_y
                    else:
                        if current_words:
                            lines.append(" ".join(current_words))
                        current_words = [token]
                        current_line_y = y

                if current_words:
                    lines.append(" ".join(current_words))

                merged = "\n".join(lines).strip()
                if merged:
                    return merged
        except Exception:
            pass

        # 4) rawdict提取span文本（最后兜底）
        try:
            raw = page.get_text("rawdict") or {}
            blocks = raw.get("blocks", []) if isinstance(raw, dict) else []
            spans: List[str] = []
            for blk in blocks:
                for line in blk.get("lines", []) if isinstance(blk, dict) else []:
                    for span in line.get("spans", []) if isinstance(line, dict) else []:
                        text_part = str(span.get("text", "")).strip()
                        if text_part:
                            spans.append(text_part)
            if spans:
                return "\n".join(spans)
        except Exception:
            pass

        return ""


class DocxParser(BaseParser):
    """DOCX文件解析器 - 增强版，保留段落结构和标题层级"""
    
    def parse(self, file_path: str) -> str:
        """解析DOCX文件并优先转换为Markdown"""
        markdown_text = self._parse_with_mammoth(file_path)
        if markdown_text:
            return markdown_text

        # 回退到原有结构化解析
        try:
            from docx import Document
            
            doc = Document(file_path)
            content_parts = []
            
            # 解析段落（保留标题层级信息）
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    # 检查是否是标题
                    if paragraph.style.name.startswith('Heading'):
                        # 标题前后添加空行，增强段落分隔
                        content_parts.append(f"\n{paragraph.text}\n")
                    else:
                        content_parts.append(paragraph.text)
            
            # 解析表格（表格前后添加明确分隔）
            for table in doc.tables:
                table_content = []
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        table_content.append(' | '.join(row_text))
                
                if table_content:
                    # 表格作为独立块，前后用双换行分隔
                    content_parts.append('\n[表格]\n' + '\n'.join(table_content))
            
            # 段落间用双换行分隔（为递归分割提供明确边界）
            plain_text = '\n\n'.join(part.strip() for part in content_parts if part.strip())
            return MarkdownNormalizer.infer_markdown_structure(plain_text)
            
        except Exception as e:
            logger.error(f"解析DOCX文件失败: {file_path}, {str(e)}")
            raise

    def _parse_with_mammoth(self, file_path: str) -> Optional[str]:
        try:
            import mammoth

            with open(file_path, "rb") as file_obj:
                result = mammoth.convert_to_markdown(file_obj)
                markdown_text = result.value
                if markdown_text and len(markdown_text.strip()) > 0:
                    logger.info(f"DOCX使用mammoth转换Markdown成功: {file_path}")
                    return MarkdownNormalizer.normalize(markdown_text)
        except Exception as error:
            logger.warning(f"mammoth不可用或转换失败，回退python-docx: {str(error)}")
        return None


class HTMLParser(BaseParser):
    """HTML文件解析器 - 按DOM顺序遍历，保留文档结构"""

    BLOCK_TAGS = frozenset([
        'h1', 'h2', 'h3', 'h4', 'h5', 'h6',
        'p', 'div', 'article', 'section', 'blockquote',
        'ul', 'ol', 'li', 'pre', 'table', 'tr',
        'header', 'main', 'aside', 'details', 'summary',
    ])
    HEADING_TAGS = frozenset(['h1', 'h2', 'h3', 'h4', 'h5', 'h6'])
    SKIP_TAGS = frozenset(['script', 'style', 'nav', 'footer', 'noscript', 'svg', 'iframe'])

    def parse(self, file_path: str) -> str:
        """按DOM原始顺序解析HTML，标题和正文保持上下文关系"""
        try:
            from bs4 import BeautifulSoup, NavigableString, Tag

            with open(file_path, 'r', encoding='utf-8') as f:
                html_content = f.read()

            soup = BeautifulSoup(html_content, 'lxml')

            for tag in soup.find_all(list(self.SKIP_TAGS)):
                tag.decompose()

            content_parts: List[str] = []
            seen_texts: set = set()

            body = soup.body or soup
            self._walk_dom(body, content_parts, seen_texts, Tag, NavigableString)

            return '\n\n'.join(part for part in content_parts if part)

        except Exception as e:
            logger.error(f"解析HTML文件失败: {file_path}, {str(e)}")
            raise

    def _walk_dom(self, element, parts: List[str], seen: set, TagClass, StringClass) -> None:
        """递归遍历DOM，按原始顺序提取文本块"""
        for child in element.children:
            if isinstance(child, StringClass):
                continue

            if not isinstance(child, TagClass):
                continue

            tag_name = (child.name or '').lower()
            if tag_name in self.SKIP_TAGS:
                continue

            if tag_name in self.HEADING_TAGS:
                text = child.get_text(separator=' ', strip=True)
                if text and text not in seen:
                    level = int(tag_name[1])
                    prefix = '#' * level
                    parts.append(f"{prefix} {text}")
                    seen.add(text)
                continue

            if tag_name in ('ul', 'ol'):
                items: List[str] = []
                for li in child.find_all('li', recursive=False):
                    li_text = li.get_text(separator=' ', strip=True)
                    if li_text:
                        items.append(f"- {li_text}")
                if items:
                    block = '\n'.join(items)
                    if block not in seen:
                        parts.append(block)
                        seen.add(block)
                continue

            if tag_name == 'table':
                rows: List[str] = []
                for tr in child.find_all('tr'):
                    cells = [td.get_text(separator=' ', strip=True) for td in tr.find_all(['td', 'th'])]
                    if any(cells):
                        rows.append(' | '.join(cells))
                if rows:
                    block = '\n'.join(rows)
                    if block not in seen:
                        parts.append(block)
                        seen.add(block)
                continue

            if tag_name in ('p', 'blockquote', 'pre'):
                text = child.get_text(separator=' ', strip=True)
                if text and text not in seen:
                    parts.append(text)
                    seen.add(text)
                continue

            if tag_name in ('div', 'article', 'section', 'main', 'aside', 'details'):
                self._walk_dom(child, parts, seen, TagClass, StringClass)
                continue

            text = child.get_text(separator=' ', strip=True)
            if text and len(text) > 20 and text not in seen:
                parts.append(text)
                seen.add(text)


class JsonParser(BaseParser):
    """JSON/JSONL文件解析器 - 将结构化数据转为可读文本"""

    def parse(self, file_path: str) -> str:
        """解析JSON/JSONL文件为结构化文本"""
        try:
            import json as _json

            with open(file_path, 'r', encoding='utf-8') as f:
                raw = f.read().strip()

            if not raw:
                return ""

            # JSONL: 逐行解析
            if file_path.endswith('.jsonl') or ('\n' in raw and not raw.startswith(('[', '{'))):
                return self._parse_jsonl(raw, _json)

            try:
                data = _json.loads(raw)
            except _json.JSONDecodeError:
                return self._parse_jsonl(raw, _json)

            return self._render(data, _json)

        except Exception as e:
            logger.error(f"解析JSON文件失败: {file_path}, {str(e)}")
            raise

    def _parse_jsonl(self, raw: str, _json) -> str:
        parts: List[str] = []
        for i, line in enumerate(raw.split('\n'), 1):
            line = line.strip()
            if not line:
                continue
            try:
                obj = _json.loads(line)
                parts.append(f"[Record {i}]\n{self._render(obj, _json)}")
            except _json.JSONDecodeError:
                parts.append(f"[Record {i}]\n{line}")
        return '\n\n'.join(parts)

    def _render(self, data, _json, indent: int = 0) -> str:
        """将JSON数据递归渲染为键值文本"""
        if isinstance(data, dict):
            lines: List[str] = []
            prefix = '  ' * indent
            for key, value in data.items():
                if isinstance(value, (dict, list)):
                    lines.append(f"{prefix}{key}:")
                    lines.append(self._render(value, _json, indent + 1))
                else:
                    lines.append(f"{prefix}{key}: {value}")
            return '\n'.join(lines)
        elif isinstance(data, list):
            if not data:
                return ""
            if all(isinstance(item, (str, int, float, bool)) for item in data):
                prefix = '  ' * indent
                return '\n'.join(f"{prefix}- {item}" for item in data)
            parts: List[str] = []
            for i, item in enumerate(data):
                parts.append(self._render(item, _json, indent))
                if i < len(data) - 1:
                    parts.append("")
            return '\n'.join(parts)
        else:
            return '  ' * indent + str(data)


class FileParser:
    """文件解析器工厂"""
    
    PARSERS = {
        '.txt': TxtParser(),
        '.md': MarkdownParser(),
        '.markdown': MarkdownParser(),
        '.pdf': PDFParser(),
        '.docx': DocxParser(),
        '.doc': DocxParser(),
        '.html': HTMLParser(),
        '.htm': HTMLParser(),
        '.json': JsonParser(),
        '.jsonl': JsonParser(),
    }
    
    @classmethod
    def parse(cls, file_path: str) -> str:
        """
        根据文件扩展名自动选择解析器
        
        Args:
            file_path: 文件路径
            
        Returns:
            解析后的文本内容
            
        Raises:
            ValueError: 不支持的文件格式
        """
        ext = os.path.splitext(file_path)[1].lower()
        
        parser = cls.PARSERS.get(ext)
        if not parser:
            raise ValueError(f"不支持的文件格式: {ext}")
        
        logger.info(f"开始解析文件: {file_path}")
        
        try:
            text = parser.parse(file_path)
            logger.info(f"文件解析成功: {file_path}, 文本长度: {len(text)}")
            return text
        except Exception as e:
            logger.error(f"文件解析失败: {file_path}, 错误: {str(e)}")
            raise
    
    @classmethod
    def is_supported(cls, filename: str) -> bool:
        """
        检查文件格式是否支持
        
        Args:
            filename: 文件名
            
        Returns:
            是否支持
        """
        ext = os.path.splitext(filename)[1].lower()
        return ext in cls.PARSERS


PARSER_MAP = {
    'txt': TxtParser,
    'md': MarkdownParser,
    'pdf': PDFParser,
    'docx': DocxParser,
    'html': HTMLParser,
    'json': JsonParser,
    'jsonl': JsonParser,
}


def get_file_parser(file_type: str) -> BaseParser:
    """
    根据文件类型获取解析器
    
    Args:
        file_type: 文件类型（不带点号，如 'txt', 'pdf'）
        
    Returns:
        解析器实例
        
    Raises:
        ValueError: 不支持的文件类型
    """
    parser_class = PARSER_MAP.get(file_type.lower())
    
    if not parser_class:
        raise ValueError(f"不支持的文件类型: {file_type}")
    
    return parser_class()
